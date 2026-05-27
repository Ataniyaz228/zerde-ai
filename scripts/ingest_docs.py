import asyncio
import os
import sys
import hashlib
import re
from pathlib import Path
import fitz  # pymupdf
import docx  # python-docx

# Add project root to sys.path so we can import zerde modules
sys.path.append(str(Path(__file__).parent.parent))

from zerde.models import EvidenceChunk, LegalRank, WebTier
from zerde.utils.cache import CacheManager
from zerde.stages.s3_gather import _regex_split_articles, _LAW_ID_KNOWN
from zerde.utils.law_registry import reload_registry

# Reverse lookup for known codes from s3_gather
_CODE_TO_LAW_ID = {v.upper(): k for k, v in _LAW_ID_KNOWN.items()}


def _extract_title_from_text(text: str, lang: str = "ru") -> str:
    """
    Извлекает заголовок закона из первых 2000 символов документа.
    
    Поддерживает два паттерна:
    - "Закон Республики Казахстан ... O ..." -> "О ..."
    - "Защите ..." (название просто)
    - "Tutulu ...", "Туралы ..."
    """
    sample = text[:2000]
    
    # Русский: "Закон Республики Казахстан ... О ..."
    m = re.search(
        r"О\s+(.{10,120}?)(?:\n|$)",
        sample,
        re.IGNORECASE,
    )
    if m:
        title = m.group(1).strip().rstrip(".")
        if len(title) > 10:
            return title
    
    # Фоллбэк: первая длинная строка (название закона)
    lines = [l.strip() for l in sample.splitlines() if l.strip()]
    for line in lines[:20]:
        if 15 < len(line) < 200 and not re.match(r"^\d", line):
            return line
    
    return ""


def detect_law_id(file_path: Path, text: str = "") -> str:
    filename = file_path.name.lower()
    parent_name = file_path.parent.name.lower()
    
    # 1. Check path/filename clues first
    # Check if this is Constitution of RK
    if "конституция" in filename or "конституция" in parent_name:
        if "k950001000" in filename:
            return "K950001000_"
        elif "k2600000000" in filename:
            return "K2600000000_"
        return "K950001000_"
        
    # Search for known full codes (like K1400000235) in filename
    for code, short_id in _CODE_TO_LAW_ID.items():
        if code.lower() in filename or code.lower().replace("_", "") in filename:
            return short_id
            
    # Fallback to extracting short ID patterns from name (e.g. 235-V)
    match = re.search(r"\b(\d+-[ivxldc]+)\b", filename, re.IGNORECASE)
    if match:
        return match.group(1).upper()
        
    # Check parent folder for code
    for code, short_id in _CODE_TO_LAW_ID.items():
        if code.lower() in parent_name or code.lower().replace("_", "") in parent_name:
            return short_id

    # 2. Check document content header if path was not enough
    if text:
        header_sample = text[:2000]
        # Check for Constitution in header text
        if "конституция" in header_sample.lower():
            if "k2600000000" in header_sample.lower():
                return "K2600000000_"
            return "K950001000_"
            
        # Check for known full codes in header text
        for code, short_id in _CODE_TO_LAW_ID.items():
            if code.lower() in header_sample.lower() or code.lower().replace("_", "") in header_sample.lower():
                return short_id

        # Check for short ID patterns in the header text (e.g. № 94-V)
        patterns = [
            r"№\s*(\d+-[IVXLCD]+)",
            r"закон\s+рк\s+от\s+.*?№\s*(\d+-[IVXLCD]+)",
            r"закон\s+республики\s+казахстан\s+.*?№\s*(\d+-[IVXLCD]+)",
            r"\b(\d+-[IVXLCD]+)\b"
        ]
        for pat in patterns:
            match = re.search(pat, header_sample, re.IGNORECASE)
            if match:
                val = match.group(1).upper()
                if val in _LAW_ID_KNOWN:
                    return val
                return val
            
    return "UNKNOWN"

def extract_pdf_text(path: Path) -> str:
    doc = fitz.open(str(path))
    pages = []
    for page in doc:
        pages.append(page.get_text("text"))
    doc.close()
    text = "\n".join(pages)
    
    # 1. Apply spaced text fixer from stage 1 to clean up Adilet PDF spacing
    from zerde.stages.s1_ingest import _fix_spaced_pdf_text
    text = _fix_spaced_pdf_text(text)
    
    # 2. Fix layout splits in article numbers: 10-\n5 -> 10-5
    text = re.sub(r"(\d+)-\s*\n\s*(\d+)", r"\1-\2", text)
    # 124\n-бап -> 124-бап
    text = re.sub(r"(\d+)\s*\n\s*-(бап|бабы|бапта|бабының|бабына)", r"\1-\2", text)
    
    return text

def extract_docx_text(path: Path) -> str:
    doc = docx.Document(str(path))
    paragraphs = []
    for para in doc.paragraphs:
        paragraphs.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            paragraphs.append(" | ".join(row_text))
    return "\n".join(paragraphs)

MAX_CHUNK_SIZE = 4000
OVERLAP = 200

def split_long_article(art_num: str, content: str) -> list[dict]:
    if len(content) <= MAX_CHUNK_SIZE:
        return [{"article_num": art_num, "content": content}]
    
    # Split by clauses (e.g. \n1., \n2., \n1), \n2))
    parts = re.split(r'\n(?=\d+[\.\)])', content)
    
    if len(parts) == 1:
        # Fallback: split by character chunks with overlap
        chunks = []
        start = 0
        chunk_idx = 1
        while start < len(content):
            end = start + MAX_CHUNK_SIZE
            chunks.append({
                "article_num": f"{art_num}_p{chunk_idx}",
                "content": content[start:end]
            })
            chunk_idx += 1
            start += MAX_CHUNK_SIZE - OVERLAP
        return chunks
        
    chunks = []
    current = ""
    chunk_idx = 1
    
    for part in parts:
        if len(current) + len(part) > MAX_CHUNK_SIZE and current:
            chunks.append({
                "article_num": f"{art_num}_p{chunk_idx}",
                "content": current
            })
            chunk_idx += 1
            current = current[-OVERLAP:] + "\n" + part
        else:
            current = (current + "\n" + part).strip()
            
    if current:
        chunks.append({
            "article_num": f"{art_num}_p{chunk_idx}",
            "content": current
        })
    return chunks

async def ingest_all_docs():
    import datetime
    print("🚀 Starting local document ingestion to zerde_cache.db...")
    cache = CacheManager("zerde_cache.db")
    docs_dir = Path("docs")
    
    all_chunks = []
    law_titles: dict[str, dict] = {}  # law_id → {title_ru, title_kz}
    
    for path in docs_dir.rglob("*"):
        if not path.is_file():
            continue
        # Automatically ignore any drafts/test bills located directly in the root of the docs/ directory
        if path.parent == docs_dir:
            print(f"Skipping test/draft file in docs root: {path.name}")
            continue
            
        suffix = path.suffix.lower()
        if suffix not in (".pdf", ".docx"):
            continue
            
        print(f"\nProcessing reference file: {path}")
        
        try:
            if suffix == ".pdf":
                text = extract_pdf_text(path)
            else:
                text = extract_docx_text(path)
                
            law_id = detect_law_id(path, text)
            print(f"Detected Law ID: {law_id}")
            
            if not text.strip():
                print(f"⚠️ Warning: Extracted text for {path} is empty or whitespace only!")
                continue
                
            print(f"Extracted {len(text)} characters. Splitting into articles...")
            
            articles = _regex_split_articles(text)
            if not articles:
                # If regex fails, split by paragraphs/pages to ensure we get chunks
                print(f"⚠️ Regex article split returned 0 articles for {path.name}. Falling back to paragraph chunks.")
                paragraphs = [p.strip() for p in text.split("\n\n")]
                valid_count = 0
                for idx, p in enumerate(paragraphs):
                    if not p:
                        print(f"   [Fallback] Paragraph {idx+1} is empty, skipping.")
                        continue
                    if len(p) < 10:
                        print(f"   [Fallback] Paragraph {idx+1} is too short ({len(p)} chars): '{p}', skipping.")
                        continue
                    articles.append({
                        "article_num": f"p{idx+1}",
                        "content": p
                    })
                    valid_count += 1
                print(f"   [Fallback] Generated {valid_count} valid paragraph chunks out of {len(paragraphs)} total paragraphs.")
                     
            # Определяем язык и заголовок документа
            filename = path.name.lower()
            
            # Language: ru or kk
            lang = "ru"
            if "kaz" in filename or "kk" in filename or ".kaz." in filename:
                lang = "kk"
            elif "rus" in filename or "ru" in filename or ".rus." in filename:
                lang = "ru"

            # Извлекаем заголовок
            doc_title = _extract_title_from_text(text, lang)
            title_ru = doc_title if lang == "ru" else ""
            title_kz = doc_title if lang == "kk" else ""

            # Сохраняем заголовок для реестра
            if law_id not in law_titles:
                law_titles[law_id] = {"title_ru": "", "title_kz": ""}
            if title_ru:
                law_titles[law_id]["title_ru"] = title_ru
            if title_kz:
                law_titles[law_id]["title_kz"] = title_kz
                
            # Source version date (format: DD-MM-YYYY -> YYYY-MM-DD)
            date_match = re.search(r"\b(\d{2})-(\d{2})-(\d{4})\b", filename)
            if date_match:
                day, month, year = date_match.groups()
                version = f"{year}-{month}-{day}"
                eff_date = datetime.date(int(year), int(month), int(day))
            else:
                version = None
                eff_date = None
                
            ingest_dt = datetime.datetime.now(datetime.timezone.utc).isoformat()
            
            # Apply adaptive chunking to each article
            adaptive_articles = []
            for art in articles:
                art_num = art["article_num"]
                content = art["content"]
                adaptive_articles.extend(split_long_article(art_num, content))
                
            print(f"Generated {len(adaptive_articles)} adaptive chunks/articles (split from {len(articles)} raw articles).")
            
            is_code = law_id.startswith("K") or "кодекс" in str(path).lower() or law_id in ("235-V", "226-V", "350-VI", "212-IV", "1000-XIII", "409-I", "442-II", "414-I")
            rank = LegalRank.CODE if is_code else LegalRank.LAW_RK
            
            for art in adaptive_articles:
                content = art["content"]
                art_num = art["article_num"]
                
                chunk_id = hashlib.sha256(content.encode()).hexdigest()
                chunk = EvidenceChunk(
                    chunk_id=chunk_id,
                    source_url=f"https://adilet.zan.kz/rus/docs/{law_id}#{art_num}",
                    source_title=f"НПА {law_id} | Ст. {art_num}",
                    content=content,
                    legal_rank=rank,
                    law_id=law_id,
                    article=art_num,
                    effective_date=eff_date,
                    language=lang,
                    source_version=version,
                    ingest_date=ingest_dt,
                )
                all_chunks.append(chunk)
                
        except Exception as e:
            print(f"❌ Error processing {path}: {e}")
            import traceback
            traceback.print_exc()
            
    if all_chunks:
        print(f"\nStoring {len(all_chunks)} chunks to sqlite cache...")
        stored = await cache.put_many(all_chunks)
        print(f"🎉 Stored {stored} chunks in zerde_cache.db successfully!")
        
        # Обновляем law_metadata для всех law_id из загруженных чанков
        law_stats: dict[str, dict] = {}
        for chunk in all_chunks:
            lid = chunk.law_id or "UNKNOWN"
            if lid not in law_stats:
                law_stats[lid] = {"chunk_count": 0, "title_ru": "", "title_kz": ""}
            law_stats[lid]["chunk_count"] += 1
            # Добавляем заголовки из накопленных за время инджеста
            if lid in law_titles:
                if not law_stats[lid]["title_ru"] and law_titles[lid].get("title_ru"):
                    law_stats[lid]["title_ru"] = law_titles[lid]["title_ru"]
                if not law_stats[lid]["title_kz"] and law_titles[lid].get("title_kz"):
                    law_stats[lid]["title_kz"] = law_titles[lid]["title_kz"]
        
        for lid, stats in law_stats.items():
            if lid == "UNKNOWN":
                continue
            # Получаем adilet_code из статического словаря
            adilet_code = _LAW_ID_KNOWN.get(lid, "")
            cache.upsert_law_metadata(
                law_id=lid,
                adilet_code=adilet_code,
                title_ru=stats.get("title_ru", ""),
                title_kz=stats.get("title_kz", ""),
                chunk_count=stats["chunk_count"],
            )
        print(f"📚 Law registry updated: {len(law_stats)} laws written to law_metadata.")
        
        # Перезагружаем синглтон реестра после инджеста
        reload_registry()
        print("✅ Law registry reloaded.")
        
        # Optimize: Only embed chunks that don't have embeddings in the database yet
        with cache._conn() as conn:
            existing_emb_ids = set(row[0] for row in conn.execute("SELECT chunk_id FROM evidence_embeddings").fetchall())
        chunks_to_embed = [c for c in all_chunks if c.chunk_id not in existing_emb_ids]
        
        if chunks_to_embed:
            print(f"\n🧠 Pre-computing BGE-M3 vector embeddings for {len(chunks_to_embed)} missing chunks...")
            # Since embed_chunks is sync, we run it directly. It will lazy-init the BGE-M3 model on GPU/CPU.
            cache.embed_chunks(chunks_to_embed)
            print(f"🎉 Successfully generated and stored BGE-M3 vector embeddings for {len(chunks_to_embed)} chunks!")
        else:
            print("\n🎉 All chunks already have vector embeddings. Skipping embedding generation!")
    else:
        print("\nNo chunks found to store!")

if __name__ == "__main__":
    asyncio.run(ingest_all_docs())

