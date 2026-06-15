"""S1 — извлечение текста из входного файла в DocumentState.

PDF читаем через pymupdf; если текстовый слой почти пустой (< 100 символов,
скан) — откатываемся на Tesseract OCR. DOCX берём через python-docx (параграфы
и таблицы), TXT — с автодетектом кодировки. Дальше нормализуем: казахскую
латиницу транслитерируем в кириллицу, язык определяем через langdetect.
"""

from __future__ import annotations

import hashlib
import logging
import os
from pathlib import Path

from zerde.models import DocumentFormat, DocumentState
from zerde.utils.kz_translit import normalize_kz_text

logger = logging.getLogger(__name__)

# Минимум символов для "нормального" PDF текстового слоя
_PDF_MIN_TEXT_CHARS = 100

# OCR конфиг — управляется через env (default: включён, rus+kaz, 200 dpi, ≤50 страниц).
# tessdata: traineddata скачиваются в ~/.local/share/tessdata
_DEFAULT_TESSDATA = os.path.expanduser("~/.local/share/tessdata")


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


async def ingest_document(file_path: str | Path) -> DocumentState:
    """
    Этап 1: Загружает файл, извлекает текст, нормализует.

    Args:
        file_path: Путь к документу.

    Returns:
        DocumentState с нормализованным текстом.
    """
    path = Path(file_path).resolve()
    logger.info(f"[S1] Ingesting: {path}")

    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    fmt = _detect_format(path)
    raw_text = await _extract_text(path, fmt)

    if not raw_text.strip():
        raise ValueError(f"No text extracted from: {path}")

    normalized = normalize_kz_text(raw_text)
    doc_id = hashlib.sha256(raw_text.encode()).hexdigest()
    language = _detect_language(normalized)

    state = DocumentState(
        doc_id=doc_id,
        original_path=str(path),
        format=fmt,
        raw_text=raw_text,
        normalized_text=normalized,
        char_count=len(normalized),
        language_detected=language,
    )

    logger.info(f"[S1] Done. doc_id={doc_id[:8]}… words={state.word_count} lang={language} fmt={fmt}")
    return state


# ---------------------------------------------------------------------------
# Format Detection
# ---------------------------------------------------------------------------


def _detect_format(path: Path) -> DocumentFormat:
    suffix = path.suffix.lower().lstrip(".")
    try:
        return DocumentFormat(suffix)
    except ValueError:
        raise ValueError(f"Unsupported format: '{suffix}'. Supported: pdf, docx, txt")


# ---------------------------------------------------------------------------
# Text Extraction
# ---------------------------------------------------------------------------


async def _extract_text(path: Path, fmt: DocumentFormat) -> str:
    match fmt:
        case DocumentFormat.TXT:
            return _extract_txt(path)
        case DocumentFormat.PDF:
            return _extract_pdf(path)
        case DocumentFormat.DOCX:
            return _extract_docx(path)
        case _:
            raise ValueError(f"Unhandled format: {fmt}")


def _extract_txt(path: Path) -> str:
    """TXT с автодетектом кодировки через chardet."""
    import chardet

    raw_bytes = path.read_bytes()
    detected = chardet.detect(raw_bytes)
    encoding = detected.get("encoding") or "utf-8"
    confidence = detected.get("confidence", 0)

    logger.debug(f"[S1/TXT] Detected encoding: {encoding} (confidence={confidence:.2f})")
    return raw_bytes.decode(encoding, errors="replace")


def _extract_pdf(path: Path) -> str:
    """
    PDF → текст через pymupdf.
    Если текстовый слой пустой (скан) — пробуем OCR через pymupdf.fitz.
    """
    import fitz  # pymupdf

    doc = fitz.open(str(path))
    pages_text: list[str] = []

    for page_num, page in enumerate(doc):
        text = page.get_text("text")  # type: ignore[arg-type]
        pages_text.append(text)
        logger.debug(f"[S1/PDF] Page {page_num + 1}: {len(text)} chars")

    full_text = "\n".join(pages_text)
    doc.close()

    # Если текстового слоя нет — используем pymupdf встроенный TextPage (blocks fallback)
    if len(full_text.strip()) < _PDF_MIN_TEXT_CHARS:
        logger.warning(f"[S1/PDF] Text layer too sparse ({len(full_text.strip())} chars). Trying blocks fallback.")
        full_text = _extract_pdf_ocr_fallback(path)

    # Если и blocks пустой — это скан, запускаем настоящий OCR (Tesseract via pymupdf)
    if len(full_text.strip()) < _PDF_MIN_TEXT_CHARS:
        logger.warning("[S1/PDF] Blocks fallback also sparse. Trying Tesseract OCR.")
        full_text = _extract_pdf_via_ocr(path)

    # Пост-процессинг: склеиваем разнесённые буквы (Adilet PDF артефакт)
    full_text = _fix_spaced_pdf_text(full_text)

    # Исправляем разрывы строк в номерах статей (10-\n5 -> 10-5, 124\n-бап -> 124-бап)
    import re
    full_text = re.sub(r"(\d+)-\s*\n\s*(\d+)", r"\1-\2", full_text)
    full_text = re.sub(r"(\d+)\s*\n\s*-(бап|бабы|бапта|бабының|бабына)", r"\1-\2", full_text)

    return full_text


def _fix_spaced_pdf_text(text: str) -> str:
    """
    Исправляет разнесённые буквы в PDF из Adilet/госорганов.
    'с т а т ь е  1 9 6' → 'статье 196'
    'З А К О Н' → 'ЗАКОН'

    Эвристика: обрабатывает построчно. Если строка содержит
    одиночные символы разделённые пробелами — склеивает их.
    """
    import re

    lines = text.split("\n")
    fixed_lines: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            fixed_lines.append(line)
            continue

        # Паттерн: последовательность одиночных символов через пробел(ы)
        # "с т а т ь е" → 6 одиночных кириллических букв
        # "1 9 6" → 3 одиночных цифры
        # Проверяем: если >40% слов — одиночные символы, строка "разнесена"
        words = stripped.split()
        single_chars = sum(1 for w in words if len(w) == 1)

        if len(words) > 2 and single_chars / len(words) > 0.4:
            # PDF из Adilet: одинарный пробел между буквами слова,
            # двойной+ пробел между словами.
            # "2 )  в  с т а т ь е  1 9 7" → ["2 )", "в", "с т а т ь е", "1 9 7"]
            # Разбиваем по 2+ пробелам → склеиваем внутри каждой группы
            word_groups = re.split(r'  +', stripped)
            result_parts: list[str] = []
            for group in word_groups:
                group = group.strip()
                if not group:
                    continue
                g_words = group.split()
                g_singles = sum(1 for w in g_words if len(w) == 1)
                if len(g_words) > 1 and g_singles / len(g_words) > 0.5:
                    # Группа разнесённых символов → склеиваем
                    result_parts.append("".join(g_words))
                else:
                    result_parts.append(group)
            result = " ".join(result_parts)
            fixed_lines.append(result)
            logger.debug(f"[S1/PDF Fix] '{stripped[:50]}' → '{result[:50]}'")
        else:
            fixed_lines.append(line)

    result = "\n".join(fixed_lines)
    if result != text:
        original_len = len(text)
        new_len = len(result)
        logger.info(f"[S1/PDF] Fixed spaced text: {original_len} → {new_len} chars ({original_len - new_len} removed)")

    return result


def _extract_pdf_ocr_fallback(path: Path) -> str:
    """
    Промежуточный fallback: pymupdf blocks-представление.
    Не настоящий OCR — просто другой способ обхода текстового слоя.
    Помогает на PDF, где обычный get_text("text") теряет блоки.
    """
    import fitz

    doc = fitz.open(str(path))
    parts: list[str] = []

    for page in doc:
        blocks = page.get_text("blocks")  # type: ignore[arg-type]
        for block in blocks:
            if len(block) >= 5 and isinstance(block[4], str):
                parts.append(block[4].strip())

    doc.close()
    result = "\n".join(p for p in parts if p)
    logger.info(f"[S1/PDF blocks fallback] Extracted {len(result)} chars from {path.name}")
    return result


def _extract_pdf_via_ocr(path: Path) -> str:
    """
    OCR через Tesseract (встроенный в pymupdf). Для сканированных PDF без
    текстового слоя.

    Управляется env-переменными (для тестов/smoke можно отключить):
      ZERDE_PDF_OCR_ENABLED  — "0" чтобы выключить (default: вкл)
      ZERDE_TESSDATA         — путь к директории с .traineddata
      ZERDE_PDF_OCR_LANGS    — "rus+kaz" (default)
      ZERDE_PDF_OCR_DPI      — 200 (default)
      ZERDE_PDF_OCR_MAX_PAGES — 50 (default; ограничение чтобы не зависнуть на 500-стр PDF)
    """
    if os.environ.get("ZERDE_PDF_OCR_ENABLED", "1").lower() in ("0", "false", "no"):
        logger.info(f"[S1/PDF OCR] disabled via ZERDE_PDF_OCR_ENABLED=0, skipping {path.name}")
        return ""

    tessdata = os.environ.get("ZERDE_TESSDATA", _DEFAULT_TESSDATA)
    if not Path(tessdata).is_dir():
        logger.error(
            f"[S1/PDF OCR] tessdata directory not found: {tessdata}. "
            f"Download rus.traineddata and kaz.traineddata into it. Skipping OCR."
        )
        return ""

    langs = os.environ.get("ZERDE_PDF_OCR_LANGS", "rus+kaz")
    dpi = int(os.environ.get("ZERDE_PDF_OCR_DPI", "200"))
    max_pages = int(os.environ.get("ZERDE_PDF_OCR_MAX_PAGES", "50"))

    import fitz

    doc = fitz.open(str(path))
    total_pages = doc.page_count
    n = min(total_pages, max_pages)
    parts: list[str] = []
    for i in range(n):
        try:
            tp = doc[i].get_textpage_ocr(
                language=langs, dpi=dpi, full=True, tessdata=tessdata
            )
            parts.append(tp.extractText())
        except Exception as e:
            logger.warning(f"[S1/PDF OCR] page {i + 1} failed for {path.name}: {e}")
    doc.close()

    if total_pages > max_pages:
        logger.warning(
            f"[S1/PDF OCR] {path.name}: OCR'd {n}/{total_pages} pages (max_pages={max_pages})"
        )

    full_text = "\n".join(p for p in parts if p)
    logger.info(f"[S1/PDF OCR] {path.name}: extracted {len(full_text)} chars from {n} pages")
    return full_text


def _extract_docx(path: Path) -> str:
    """DOCX → текст: параграфы + таблицы через python-docx."""
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []

    # Параграфы
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Таблицы
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Language Detection
# ---------------------------------------------------------------------------


def _detect_language(text: str) -> str:
    """
    Определяет язык через langdetect.
    Fallback: эвристика по кириллице если langdetect упадёт.
    """
    # 0. Эвристика по уникальным казахским кириллическим буквам
    sample = text[:3000]
    sample_lower = sample.lower()
    kk_cyrillic_chars = set("әіңғүұқөһ")
    if any(c in kk_cyrillic_chars for c in sample_lower):
        return "kk"

    try:
        from langdetect import DetectorFactory, detect
        DetectorFactory.seed = 0  # Детерминированность

        sample = text[:3000]  # Достаточно для детектирования
        lang = detect(sample)

        # H6 Fix: Проверяем kaz/kk первыми, чтобы kaz не уходило в дефолтный ru
        if lang in ("kk", "kaz"):
            return "kk"
        elif lang in ("ru", "en"):
            return lang
        else:
            return "ru"  # Для КЗ юридических документов — дефолт ru
    except Exception as e:
        logger.debug(f"[S1] langdetect failed: {e}. Using heuristic.")
        return _heuristic_lang(text)


def _heuristic_lang(text: str) -> str:
    """Эвристика: считаем кириллицу vs латиницу с поддержкой казахской латиницы."""
    cyrillic = sum(1 for c in text[:500] if "\u0400" <= c <= "\u04ff")
    latin = sum(1 for c in text[:500] if "a" <= c.lower() <= "z" or c.lower() in "áñóúığşý")

    if cyrillic > latin * 2:
        return "ru"
    elif latin > cyrillic * 2:
        sample_lower = text[:1000].lower()
        kk_latin_chars = set("áñóúığşý")
        has_kk_chars = any(c in kk_latin_chars for c in sample_lower)

        kk_latin_words = {"jane", "jáne", "bolady", "qazaqstan", "turaly", "týraly", "kodeksi", "zań", "zang", "memleket", "respyblıkasy", "respublika"}
        words_in_sample = set(sample_lower.split())
        has_kk_words = any(w in kk_latin_words for w in words_in_sample)

        if has_kk_chars or has_kk_words:
            return "kk"
        return "en"
    else:
        return "mixed"
